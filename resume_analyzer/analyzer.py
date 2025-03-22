# import os
# import pymupdf as fitz  # PyMuPDF
# import PyPDF2
# import docx 
# import re
# from pathlib import Path
# import config

# class ResumeParser:
#     def __init__(self):
#         pass
        
#     def extract_text_from_pdf(self, file_path):
#         """Extract text from PDF file using PyMuPDF"""
#         try:
#             text = ""
#             # Open the PDF
#             with fitz.open(file_path) as doc:
#                 # Iterate through pages
#                 for page in doc:
#                     text += page.get_text()
#             return text
#         except Exception as e:
#             print(f"Error extracting text from PDF: {e}")
#             # Fallback to PyPDF2
#             try:
#                 text = ""
#                 with open(file_path, 'rb') as file:
#                     pdf_reader = PyPDF2.PdfReader(file)
#                     for page_num in range(len(pdf_reader.pages)):
#                         text += pdf_reader.pages[page_num].extract_text()
#                 return text
#             except Exception as e2:
#                 print(f"Fallback PDF extraction failed: {e2}")
#                 return ""
    
#     def extract_text_from_docx(self, file_path):
#         """Extract text from DOCX file"""
#         try:
#             doc = docx.Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs])
#             return text
#         except Exception as e:
#             print(f"Error extracting text from DOCX: {e}")
#             return ""
    
#     def extract_text_from_txt(self, file_path):
#         """Extract text from TXT file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except UnicodeDecodeError:
#             # Try different encoding
#             try:
#                 with open(file_path, 'r', encoding='latin-1') as file:
#                     return file.read()
#             except Exception as e:
#                 print(f"Error reading TXT file: {e}")
#                 return ""
#         except Exception as e:
#             print(f"Error extracting text from TXT: {e}")
#             return ""
    
#     def extract_text(self, file_path):
#         """Extract text from file based on extension"""
#         ext = Path(file_path).suffix.lower()
        
#         if ext == '.pdf':
#             return self.extract_text_from_pdf(file_path)
#         elif ext == '.docx':
#             return self.extract_text_from_docx(file_path)
#         elif ext == '.txt':
#             return self.extract_text_from_txt(file_path)
#         else:
#             print(f"Unsupported file format: {ext}")
#             return ""


# class ResumeAnalyzer:
#     def __init__(self, retriever):
#         self.parser = ResumeParser()
#         self.retriever = retriever
        
#     def analyze_resume(self, file_path, role_id):
#         """Analyze a resume for a specific role"""
#         # Extract text from resume
#         resume_text = self.parser.extract_text(file_path)
        
#         if not resume_text:
#             return {
#                 "success": False,
#                 "error": "Could not extract text from resume",
#                 "score": 0.0,
#                 "text": ""
#             }
        
#         # Calculate similarity score
#         similarity_score = self.retriever.calculate_resume_similarity(resume_text, role_id)
        
#         return {
#             "success": True,
#             "score": similarity_score,
#             "text": resume_text
#         }

import os
import pymupdf as fitz  # PyMuPDF
import PyPDF2
import docx 
import re
from pathlib import Path
import config

class ResumeParser:
    def __init__(self, retriever=None, config=None):
        self.retriever = retriever
        self.config = config
        
    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file using PyMuPDF"""
        try:
            text = ""
            # Open the PDF
            with fitz.open(file_path) as doc:
                # Iterate through pages
                for page in doc:
                    text += page.get_text()
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            # Fallback to PyPDF2
            try:
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        text += pdf_reader.pages[page_num].extract_text()
                return text
            except Exception as e2:
                print(f"Fallback PDF extraction failed: {e2}")
                return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"Error reading TXT file: {e}")
                return ""
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""
    
    def extract_text(self, file_path):
        """Extract text from file based on extension"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            print(f"Unsupported file format: {ext}")
            return ""
        
class ResumeAnalyzer:
    def __init__(self, retriever):
        self.parser = ResumeParser(retriever=retriever)
        self.retriever = retriever
        
    def analyze_resume(self, file_path, role_id):
        """Analyze a resume for a specific role"""
        # Extract text from resume
        resume_text = self.parser.extract_text(file_path)
        
        if not resume_text:
            return {
                "success": False,
                "error": "Could not extract text from resume",
                "score": 0.0,
                "text": ""
            }
        
        # Calculate similarity score
        similarity_score = self.retriever.calculate_resume_similarity(resume_text, role_id)
        
        return {
            "success": True,
            "score": similarity_score,
            "text": resume_text
        }